import os
import datetime
import json
from flask import (
    Flask,
    render_template,
    request,
    redirect,
    url_for,
    send_from_directory,
    send_file,
    flash,
)
import re
from bs4 import BeautifulSoup
from docx import Document
import bleach
from docx.shared import Inches

app = Flask(__name__)
app.secret_key = 'change-this'

# Application version
VERSION = "0.8.2"
app.jinja_env.globals['app_version'] = VERSION

DATA_DIR = os.environ.get('DATA_DIR', os.path.join(os.getcwd(), 'data'))

# Ensure data directory exists
os.makedirs(DATA_DIR, exist_ok=True)

SETTINGS_FILE = os.path.join(DATA_DIR, 'settings.json')
OPEN_BOOKS_FILE = os.path.join(DATA_DIR, 'open_books.json')
CLOSED_FOLDERS_FILE = os.path.join(DATA_DIR, 'closed_folders.json')
CLOSED_CHAPTERS_FILE = os.path.join(DATA_DIR, 'closed_chapters.json')


def load_open_books():
    if os.path.isfile(OPEN_BOOKS_FILE):
        try:
            with open(OPEN_BOOKS_FILE) as f:
                return json.load(f)
        except json.JSONDecodeError:
            pass
    books = list_all_books()
    save_open_books(books)
    return books


def save_open_books(books: list) -> None:
    with open(OPEN_BOOKS_FILE, 'w') as f:
        json.dump(books, f)


def load_closed_folders() -> list:
    if os.path.isfile(CLOSED_FOLDERS_FILE):
        try:
            with open(CLOSED_FOLDERS_FILE) as f:
                return json.load(f)
        except json.JSONDecodeError:
            pass
    return []


def save_closed_folders(folders: list) -> None:
    with open(CLOSED_FOLDERS_FILE, 'w') as f:
        json.dump(folders, f)


def load_closed_chapters() -> list:
    if os.path.isfile(CLOSED_CHAPTERS_FILE):
        try:
            with open(CLOSED_CHAPTERS_FILE) as f:
                return json.load(f)
        except json.JSONDecodeError:
            pass
    return []


def save_closed_chapters(chapters: list) -> None:
    with open(CLOSED_CHAPTERS_FILE, 'w') as f:
        json.dump(chapters, f)


def load_settings():
    defaults = {
        'dark_mode': False,
        'sidebar_color': '#f0f0f0',
        'text_color': '#000000',
        'bg_color': '#ffffff',
        'toolbar_color': '#dddddd',
        'editor_color': '#ffffff',
    }
    if os.path.isfile(SETTINGS_FILE):
        with open(SETTINGS_FILE) as f:
            try:
                data = json.load(f)
                defaults.update(data)
                return defaults
            except json.JSONDecodeError:
                pass
    return defaults


def save_settings(data: dict) -> None:
    with open(SETTINGS_FILE, 'w') as f:
        json.dump(data, f)


def safe_name(name: str) -> str:
    """Return a sanitized filename allowing common punctuation."""
    allowed = {' ', '_', '-', '.', '(', ')'}
    return ''.join(c for c in name if c.isalnum() or c in allowed).rstrip()


def html_to_text(html: str) -> str:
    """Convert HTML to plain text."""
    soup = BeautifulSoup(html, "html.parser")
    return soup.get_text(separator="\n")


def sanitize_html(html: str) -> str:
    """Strip unwanted tags to prevent script injection."""
    allowed_tags = [
        "b", "strong", "i", "em", "u", "p", "br", "div", "ul", "ol", "li", "hr",
        "span", "img"
    ]
    allowed_attrs = {
        "*": ["class", "style"],
        "img": ["src", "alt", "width", "height", "style"]
    }
    return bleach.clean(html, tags=allowed_tags, attributes=allowed_attrs, strip=True)


def html_to_docx(html: str, path: str) -> None:
    """Save limited HTML content to a DOCX file."""
    doc = Document()
    soup = BeautifulSoup(html, "html.parser")

    def process(elem, paragraph, formatting=None):
        if formatting is None:
            formatting = {}
        if isinstance(elem, str):
            run = paragraph.add_run(elem)
            run.bold = formatting.get("bold", False)
            run.italic = formatting.get("italic", False)
            run.underline = formatting.get("underline", False)
            return
        tag = elem.name
        fmt = formatting.copy()
        if tag in ("strong", "b"):
            fmt["bold"] = True
        if tag in ("em", "i"):
            fmt["italic"] = True
        if tag == "u":
            fmt["underline"] = True
        if tag == "img":
            src = elem.get("src", "")
            if src.startswith("data:image/"):
                import base64
                from io import BytesIO
                header, b64 = src.split(",", 1)
                data = base64.b64decode(b64)
                bio = BytesIO(data)
                width = elem.get("width")
                height = elem.get("height")
                w = Inches(int(width)/96) if width and width.isdigit() else None
                h = Inches(int(height)/96) if height and height.isdigit() else None
                doc.add_picture(bio, width=w, height=h)
            return
        if tag in ("p", "div", "br"):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            for child in elem.children:
                process(child, p, fmt)
            return
        for child in elem.children:
            process(child, paragraph, fmt)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    for child in soup.children:
        process(child, p)
    doc.save(path)


def append_html_to_docx(doc: Document, html: str) -> None:
    """Append HTML content to an existing DOCX document."""
    soup = BeautifulSoup(html, "html.parser")

    def process(elem, paragraph, formatting=None):
        if formatting is None:
            formatting = {}
        if isinstance(elem, str):
            run = paragraph.add_run(elem)
            run.bold = formatting.get("bold", False)
            run.italic = formatting.get("italic", False)
            run.underline = formatting.get("underline", False)
            return
        tag = elem.name
        fmt = formatting.copy()
        if tag in ("strong", "b"):
            fmt["bold"] = True
        if tag in ("em", "i"):
            fmt["italic"] = True
        if tag == "u":
            fmt["underline"] = True
        if tag == "img":
            src = elem.get("src", "")
            if src.startswith("data:image/"):
                import base64
                from io import BytesIO
                header, b64 = src.split(",", 1)
                data = base64.b64decode(b64)
                bio = BytesIO(data)
                width = elem.get("width")
                height = elem.get("height")
                w = Inches(int(width)/96) if width and width.isdigit() else None
                h = Inches(int(height)/96) if height and height.isdigit() else None
                doc.add_picture(bio, width=w, height=h)
            return
        if tag in ("p", "div", "br"):
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.5)
            for child in elem.children:
                process(child, p, fmt)
            return
        for child in elem.children:
            process(child, paragraph, fmt)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    for child in soup.children:
        process(child, p)


def sanitize_path(folder: str) -> str:
    parts = [safe_name(p) for p in folder.split('/') if p]
    return os.path.join(*parts) if parts else ''


def load_order(folder: str) -> dict:
    """Load ordering info for a folder."""
    order_file = os.path.join(DATA_DIR, sanitize_path(folder), 'order.json')
    if os.path.isfile(order_file):
        try:
            with open(order_file) as f:
                data = json.load(f)
                if 'folders' not in data:
                    data['folders'] = []
                if 'chapters' not in data:
                    data['chapters'] = []
                return data
        except json.JSONDecodeError:
            pass
    return {'folders': [], 'chapters': []}


def save_order(folder: str, order: dict) -> None:
    os.makedirs(os.path.join(DATA_DIR, sanitize_path(folder)), exist_ok=True)
    order_file = os.path.join(DATA_DIR, sanitize_path(folder), 'order.json')
    with open(order_file, 'w') as f:
        json.dump(order, f)


def list_chapters(folder: str, include_closed: bool = False):
    path = os.path.join(DATA_DIR, sanitize_path(folder))
    if not os.path.isdir(path):
        return []
    chapters = [
        c
        for c in os.listdir(path)
        if os.path.isdir(os.path.join(path, c))
        and os.path.isfile(os.path.join(path, c, 'chapter.html'))
    ]
    order = load_order(folder).get('chapters', [])
    ordered = [c for c in order if c in chapters]
    remaining = [c for c in chapters if c not in ordered]
    remaining.sort(key=lambda n: os.path.getctime(os.path.join(path, n)))
    chapters = ordered + remaining
    if not include_closed:
        closed = load_closed_chapters()
        prefix = sanitize_path(folder)
        chapters = [c for c in chapters if f"{prefix}/{c}" not in closed]
    return chapters



def list_subfolders(folder: str, include_closed: bool = False):
    path = os.path.join(DATA_DIR, sanitize_path(folder))
    if not os.path.isdir(path):
        return []
    subs = [
        c
        for c in os.listdir(path)
        if os.path.isdir(os.path.join(path, c))
        and not os.path.isfile(os.path.join(path, c, 'chapter.html'))
    ]
    order = load_order(folder).get('folders', [])
    ordered = [s for s in order if s in subs]
    remaining = [s for s in subs if s not in ordered]
    remaining.sort(key=lambda n: os.path.getctime(os.path.join(path, n)))
    subs = ordered + remaining
    if not include_closed:
        closed = load_closed_folders()
        prefix = sanitize_path(folder)
        subs = [s for s in subs if os.path.join(prefix, s) not in closed]
    return subs


def list_all_books():
    path = DATA_DIR
    books = [f for f in os.listdir(path) if os.path.isdir(os.path.join(path, f))]
    order = load_order('').get('folders', [])
    ordered = [b for b in order if b in books]
    remaining = [b for b in books if b not in ordered]
    remaining.sort(key=lambda n: os.path.getctime(os.path.join(path, n)))
    return ordered + remaining


def list_books():
    all_books = list_all_books()
    open_books = load_open_books()
    return [b for b in all_books if b in open_books]


def note_filename(chapter: str) -> str:
    """Return the standard notes filename for a chapter."""
    return f"{chapter.replace(' ', '_')}_notes.txt"


def list_notes(folder: str, chapter: str):
    """Return the notes filename for the chapter if it exists."""
    path = os.path.join(DATA_DIR, sanitize_path(folder), safe_name(chapter))
    filename = note_filename(chapter)
    note_path = os.path.join(path, filename)
    if os.path.isfile(note_path):
        return [filename]
    return []


app.jinja_env.globals['list_chapters'] = list_chapters
app.jinja_env.globals['list_notes'] = list_notes
app.jinja_env.globals['list_subfolders'] = list_subfolders
app.jinja_env.globals['list_books'] = list_books
app.jinja_env.globals['list_all_books'] = list_all_books


@app.context_processor
def inject_app_settings():
    colors = {b: read_color(b) for b in list_all_books()}
    return {'app_settings': load_settings(), 'book_colors': colors}


def read_description(folder: str) -> str:
    """Return description text for a folder if present."""
    path = os.path.join(DATA_DIR, sanitize_path(folder), 'description.txt')
    if os.path.isfile(path):
        with open(path) as f:
            return f.read()
    return ''


def write_description(folder: str, text: str) -> None:
    os.makedirs(os.path.join(DATA_DIR, sanitize_path(folder)), exist_ok=True)
    path = os.path.join(DATA_DIR, sanitize_path(folder), 'description.txt')
    with open(path, 'w') as f:
        f.write(text)

def read_author(folder: str) -> str:
    """Return author text for a folder if present."""
    path = os.path.join(DATA_DIR, sanitize_path(folder), 'author.txt')
    if os.path.isfile(path):
        with open(path) as f:
            return f.read()
    return ''

def write_author(folder: str, text: str) -> None:
    os.makedirs(os.path.join(DATA_DIR, sanitize_path(folder)), exist_ok=True)
    path = os.path.join(DATA_DIR, sanitize_path(folder), 'author.txt')
    with open(path, 'w') as f:
        f.write(text)

def read_color(folder: str) -> str:
    """Return stored color for a book if set."""
    path = os.path.join(DATA_DIR, sanitize_path(folder), 'color.txt')
    if os.path.isfile(path):
        with open(path) as f:
            return f.read().strip()
    return ''


def write_color(folder: str, color: str) -> None:
    os.makedirs(os.path.join(DATA_DIR, sanitize_path(folder)), exist_ok=True)
    path = os.path.join(DATA_DIR, sanitize_path(folder), 'color.txt')
    with open(path, 'w') as f:
        f.write(color)


@app.route('/')
def index():
    all_books = list_all_books()
    open_books = load_open_books()
    folders = [b for b in all_books if b in open_books]
    return render_template(
        'index.html',
        folders=folders,
        all_books=all_books,
        open_books=open_books,
    )


@app.route('/settings', methods=['GET', 'POST'])
def app_settings_page():
    settings = load_settings()
    if request.method == 'POST':
        if 'reset' in request.form:
            settings = {
                'dark_mode': False,
                'sidebar_color': '#f0f0f0',
                'text_color': '#000000',
                'bg_color': '#ffffff',
                'toolbar_color': '#dddddd',
                'editor_color': '#ffffff',
            }
        else:
            settings['dark_mode'] = bool(request.form.get('dark_mode'))
            settings['sidebar_color'] = request.form.get('sidebar_color', '#f0f0f0') or '#f0f0f0'
            settings['text_color'] = request.form.get('text_color', '#000000') or '#000000'
            settings['bg_color'] = request.form.get('bg_color', '#ffffff') or '#ffffff'
            settings['toolbar_color'] = request.form.get('toolbar_color', '#dddddd') or '#dddddd'
            settings['editor_color'] = request.form.get('editor_color', '#ffffff') or '#ffffff'
        save_settings(settings)
        flash('Settings saved')
        return redirect(url_for('index'))
    folders = list_books()
    return render_template('app_settings.html', settings=settings, folders=folders)


@app.route('/folder/create', methods=['POST'])
def create_folder():
    name = safe_name(request.form.get('name', ''))
    if not name:
        flash('Folder name required')
        return redirect(url_for('index'))
    path = os.path.join(DATA_DIR, name)
    os.makedirs(path, exist_ok=True)
    order = load_order('')
    if name not in order.get('folders', []):
        order.setdefault('folders', []).append(name)
        save_order('', order)
    open_books = load_open_books()
    if name not in open_books:
        open_books.append(name)
        save_open_books(open_books)
    return redirect(url_for('view_folder', folder=name))


@app.route('/wizard/book', methods=['GET', 'POST'])
def book_wizard():
    """Simple wizard to create a book with common sub-folders."""
    if request.method == 'POST':
        title = safe_name(request.form.get('title', '')).strip()
        chapters = safe_name(request.form.get('chapters', 'Chapters')).strip()
        author_text = request.form.get('author', '').strip()
        color_value = request.form.get('color', '#dddddd').strip()
        extras = request.form.getlist('extras')
        if not title:
            flash('Book title required')
            return redirect(url_for('book_wizard'))
        path = os.path.join(DATA_DIR, title)
        os.makedirs(path, exist_ok=True)
        if author_text:
            write_author(title, author_text)
        if color_value:
            write_color(title, color_value)
        created = []
        if chapters:
            chap_folder = os.path.join(path, chapters)
            os.makedirs(chap_folder, exist_ok=True)
            created.append(chapters)
            # create first chapter
            ch1 = os.path.join(chap_folder, 'Chapter One')
            os.makedirs(ch1, exist_ok=True)
            open(os.path.join(ch1, 'chapter.html'), 'a').close()
            ch_order = load_order(f"{title}/{chapters}")
            if 'Chapter One' not in ch_order.get('chapters', []):
                ch_order.setdefault('chapters', []).append('Chapter One')
                save_order(f"{title}/{chapters}", ch_order)
        for sub in extras:
            sub_name = safe_name(sub)
            if sub_name:
                os.makedirs(os.path.join(path, sub_name), exist_ok=True)
                created.append(sub_name)
        root_order = load_order('')
        if title not in root_order.get('folders', []):
            root_order.setdefault('folders', []).append(title)
            save_order('', root_order)
        order = load_order(title)
        for sub in created:
            if sub not in order.get('folders', []):
                order.setdefault('folders', []).append(sub)
        if order.get('folders'):
            save_order(title, order)
        open_books = load_open_books()
        if title not in open_books:
            open_books.append(title)
            save_open_books(open_books)
        return redirect(url_for('view_folder', folder=title))
    folders = list_books()
    return render_template('book_wizard.html', folders=folders)


@app.route('/folder/<path:folder>/delete', methods=['POST'])
def delete_folder(folder):
    folder_name = sanitize_path(folder)
    path = os.path.join(DATA_DIR, folder_name)
    parent = os.path.dirname(folder_name)
    if os.path.isdir(path):
        import shutil
        shutil.rmtree(path)
        flash('Book deleted')
        if parent:
            order = load_order(parent)
            if os.path.basename(folder_name) in order.get('folders', []):
                order['folders'].remove(os.path.basename(folder_name))
                save_order(parent, order)
    else:
        flash('Book not found')
    if parent:
        return redirect(url_for('view_folder', folder=parent))
    else:
        order = load_order('')
        bname = os.path.basename(folder_name)
        if bname in order.get('folders', []):
            order['folders'].remove(bname)
            save_order('', order)
        open_books = load_open_books()
        if bname in open_books:
            open_books.remove(bname)
            save_open_books(open_books)
    return redirect(url_for('index'))


@app.route('/folder/<path:folder>/close', methods=['POST'])
def close_folder(folder):
    """Hide a book or sub-folder from the sidebar."""
    folder_name = sanitize_path(folder)
    parent = os.path.dirname(folder_name)
    if parent:
        closed = load_closed_folders()
        if folder_name not in closed:
            closed.append(folder_name)
            save_closed_folders(closed)
        flash('Sub-folder closed')
        return redirect(url_for('view_folder', folder=parent))
    else:
        open_books = load_open_books()
        if folder_name in open_books:
            open_books.remove(folder_name)
            save_open_books(open_books)
        flash('Book closed')
        return redirect(url_for('index'))


@app.route('/folder/<path:folder>/open', methods=['POST'])
def open_folder(folder):
    """Show a previously closed book or sub-folder in the sidebar."""
    folder_name = sanitize_path(folder)
    parent = os.path.dirname(folder_name)
    if parent:
        closed = load_closed_folders()
        if folder_name in closed:
            closed.remove(folder_name)
            save_closed_folders(closed)
        flash('Sub-folder opened')
        return redirect(url_for('view_folder', folder=parent))
    else:
        open_books = load_open_books()
        if folder_name not in open_books:
            open_books.append(folder_name)
            save_open_books(open_books)
        flash('Book opened')
        return redirect(url_for('view_folder', folder=folder_name))


@app.route('/books/reorder', methods=['POST'])
def reorder_books():
    """Reorder top level books."""
    order = load_order('')
    if request.is_json:
        items = request.json.get('order', [])
        order['folders'] = items
        save_order('', order)
        return ('', 204)
    name = request.form.get('item_name')
    direction = request.form.get('direction')
    items = order.get('folders', [])
    if name in items:
        idx = items.index(name)
        if direction == 'up' and idx > 0:
            items[idx], items[idx-1] = items[idx-1], items[idx]
        elif direction == 'down' and idx < len(items)-1:
            items[idx], items[idx+1] = items[idx+1], items[idx]
        order['folders'] = items
        save_order('', order)
    return redirect(url_for('index'))


@app.route('/folder/<path:folder>/settings', methods=['GET', 'POST'])
def folder_settings(folder):
    folder_name = sanitize_path(folder)
    path = os.path.join(DATA_DIR, folder_name)
    if not os.path.isdir(path):
        flash('Book not found')
        return redirect(url_for('index'))
    description = read_description(folder_name)
    author = read_author(folder_name)
    color = read_color(folder_name)
    parent = os.path.dirname(folder_name)
    order = load_order(folder_name)
    if request.method == 'POST':
        if 'item_type' in request.form:
            typ = request.form['item_type']
            name = request.form['item_name']
            direction = request.form['direction']
            items = order.get(f'{typ}s', [])
            if name in items:
                idx = items.index(name)
                if direction == 'up' and idx > 0:
                    items[idx], items[idx-1] = items[idx-1], items[idx]
                elif direction == 'down' and idx < len(items)-1:
                    items[idx], items[idx+1] = items[idx+1], items[idx]
                order[f'{typ}s'] = items
                save_order(folder_name, order)
            return redirect(url_for('folder_settings', folder=folder_name))
        new_name = safe_name(request.form.get('name', folder_name.split('/')[-1]))
        desc = request.form.get('description', '')
        author_text = request.form.get('author', '')
        color_value = request.form.get('color', color)
        if new_name and new_name != folder_name.split('/')[-1]:
            new_path = os.path.join(DATA_DIR, os.path.dirname(folder_name), new_name)
            if os.path.exists(new_path):
                flash('Name already exists')
            else:
                os.rename(path, new_path)
                parent = os.path.dirname(folder_name)
                if parent:
                    parent_order = load_order(parent)
                    old = folder_name.split('/')[-1]
                    if old in parent_order.get('folders', []):
                        idx = parent_order['folders'].index(old)
                        parent_order['folders'][idx] = new_name
                        save_order(parent, parent_order)
                else:
                    root_order = load_order('')
                    old = folder_name.split('/')[-1]
                    if old in root_order.get('folders', []):
                        idx = root_order['folders'].index(old)
                        root_order['folders'][idx] = new_name
                        save_order('', root_order)
                    open_books = load_open_books()
                    if old in open_books:
                        open_books[open_books.index(old)] = new_name
                        save_open_books(open_books)
                folder_name = os.path.join(os.path.dirname(folder_name), new_name).strip('/')
                path = new_path
                flash('Book renamed')
        write_description(folder_name, desc)
        write_author(folder_name, author_text)
        if not parent:
            write_color(folder_name, color_value)
        return redirect(url_for('view_folder', folder=folder_name))
    subfolders = list_subfolders(folder_name, include_closed=True)
    chapters = list_chapters(folder_name, include_closed=True)
    folders = list_books()
    return render_template(
        'settings.html',
        folder=folder_name,
        name=folder_name.split('/')[-1],
        description=description,
        author=author,
        subfolders=subfolders,
        chapters=chapters,
        folders=folders,
        color=color,
    )


@app.route('/folder/<path:folder>/reorder', methods=['POST'])
def reorder_folder(folder):
    folder_name = sanitize_path(folder)
    order = load_order(folder_name)
    if request.is_json:
        typ = request.json.get('type')
        items = request.json.get('order', [])
        if typ in ('folder', 'chapter'):
            order[f'{typ}s'] = items
            save_order(folder_name, order)
        return ('', 204)
    return redirect(url_for('folder_settings', folder=folder_name))


@app.route('/folder/<path:folder>')
def view_folder(folder):
    folder_name = sanitize_path(folder)
    path = os.path.join(DATA_DIR, folder_name)
    if not os.path.isdir(path):
        flash('Folder not found')
        return redirect(url_for('index'))
    chapters = list_chapters(folder_name, include_closed=True)
    subfolders = list_subfolders(folder_name, include_closed=True)
    closed_chapters = [c for c in chapters if f"{folder_name}/{c}" in load_closed_chapters()]
    closed_subfolders = [s for s in subfolders if os.path.join(folder_name, s) in load_closed_folders()]
    open_chapters = [c for c in chapters if c not in closed_chapters]
    open_subfolders = [s for s in subfolders if s not in closed_subfolders]
    chapters = open_chapters
    subfolders = open_subfolders
    description = read_description(folder_name)
    author = read_author(folder_name)
    folders = list_books()
    return render_template('folder.html', folder=folder_name, chapters=chapters, subfolders=subfolders, folders=folders, description=description, author=author, closed_chapters=closed_chapters, closed_subfolders=closed_subfolders)


@app.route('/folder/<path:folder>/chapter/create', methods=['POST'])
def create_chapter(folder):
    folder_name = sanitize_path(folder)
    chapter = safe_name(request.form.get('name', ''))
    if not chapter:
        flash('Chapter name required')
        return redirect(url_for('view_folder', folder=folder_name))
    path = os.path.join(DATA_DIR, folder_name, chapter)
    os.makedirs(path, exist_ok=True)
    open(os.path.join(path, 'chapter.html'), 'a').close()
    order = load_order(folder_name)
    if chapter not in order.get('chapters', []):
        order.setdefault('chapters', []).append(chapter)
        save_order(folder_name, order)
    return redirect(url_for('view_chapter', folder=folder_name, chapter=chapter))




@app.route('/folder/<path:folder>/chapter/<chapter>')
def view_chapter(folder, chapter):
    folder_name = sanitize_path(folder)
    chapter_name = safe_name(chapter)
    path = os.path.join(DATA_DIR, folder_name, chapter_name)
    if not os.path.isdir(path):
        flash('Chapter not found')
        return redirect(url_for('view_folder', folder=folder_name))
    chapter_file = os.path.join(path, 'chapter.html')
    chapter_html = ''
    if os.path.isfile(chapter_file):
        with open(chapter_file) as f:
            chapter_html = f.read()

    notes_file = os.path.join(path, note_filename(chapter_name))
    notes_text = ''
    if os.path.isfile(notes_file):
        with open(notes_file) as f:
            notes_text = f.read()

    folders = list_books()
    chapters = list_chapters(folder_name, include_closed=True)
    return render_template(
        'chapter.html',
        folder=folder_name,
        chapter=chapter_name,
        notes_text=notes_text,
        folders=folders,
        chapters=chapters,
        chapter_html=chapter_html
    )





@app.route('/folder/<path:folder>/chapter/<chapter>/notes/save', methods=['POST'])
def save_notes(folder, chapter):
    folder_name = sanitize_path(folder)
    chapter_name = safe_name(chapter)
    text = request.form.get('notes', '')
    path = os.path.join(DATA_DIR, folder_name, chapter_name)
    os.makedirs(path, exist_ok=True)
    note_path = os.path.join(path, note_filename(chapter_name))
    with open(note_path, 'w') as f:
        f.write(text)
    return ('', 204)


@app.route('/folder/<path:folder>/chapter/<chapter>/save', methods=['POST'])
def save_chapter(folder, chapter):
    folder_name = sanitize_path(folder)
    chapter_name = safe_name(chapter)
    text = request.form.get('text', '')
    text = sanitize_html(text)
    path = os.path.join(DATA_DIR, folder_name, chapter_name)
    os.makedirs(path, exist_ok=True)
    html_path = os.path.join(path, 'chapter.html')
    with open(html_path, 'w') as f:
        f.write(text)
    docx_path = os.path.join(path, 'chapter.docx')
    html_to_docx(text, docx_path)
    return redirect(url_for('view_chapter', folder=folder_name, chapter=chapter_name))


@app.route('/folder/<path:folder>/chapter/<chapter>/autosave', methods=['POST'])
def autosave_chapter(folder, chapter):
    folder_name = sanitize_path(folder)
    chapter_name = safe_name(chapter)
    text = request.form.get('text', '')
    text = sanitize_html(text)
    path = os.path.join(DATA_DIR, folder_name, chapter_name)
    os.makedirs(path, exist_ok=True)
    html_path = os.path.join(path, 'chapter.html')
    with open(html_path, 'w') as f:
        f.write(text)
    docx_path = os.path.join(path, 'chapter.docx')
    html_to_docx(text, docx_path)
    return ('', 204)


@app.route('/folder/<path:folder>/chapter/<chapter>/delete', methods=['POST'])
def delete_chapter(folder, chapter):
    folder_name = sanitize_path(folder)
    chapter_name = safe_name(chapter)
    path = os.path.join(DATA_DIR, folder_name, chapter_name)
    if os.path.isdir(path):
        import shutil
        shutil.rmtree(path)
        flash('Chapter deleted')
        order = load_order(folder_name)
        if chapter_name in order.get('chapters', []):
            order['chapters'].remove(chapter_name)
            save_order(folder_name, order)
    else:
        flash('Chapter not found')
    return redirect(url_for('view_folder', folder=folder_name))


@app.route('/folder/<path:folder>/chapter/<chapter>/close', methods=['POST'])
def close_chapter(folder, chapter):
    """Hide a chapter from the sidebar."""
    folder_name = sanitize_path(folder)
    chapter_name = safe_name(chapter)
    closed = load_closed_chapters()
    key = f"{folder_name}/{chapter_name}"
    if key not in closed:
        closed.append(key)
        save_closed_chapters(closed)
    flash('Chapter closed')
    return redirect(url_for('view_folder', folder=folder_name))


@app.route('/folder/<path:folder>/chapter/<chapter>/open', methods=['POST'])
def open_chapter(folder, chapter):
    """Show a previously closed chapter."""
    folder_name = sanitize_path(folder)
    chapter_name = safe_name(chapter)
    closed = load_closed_chapters()
    key = f"{folder_name}/{chapter_name}"
    if key in closed:
        closed.remove(key)
        save_closed_chapters(closed)
    flash('Chapter opened')
    return redirect(url_for('view_folder', folder=folder_name))


@app.route('/folder/<path:folder>/chapter/<chapter>/rename', methods=['POST'])
def rename_chapter(folder, chapter):
    folder_name = sanitize_path(folder)
    chapter_name = safe_name(chapter)
    new_name = safe_name(request.form.get('new_name', ''))
    if not new_name:
        flash('New name required')
        return redirect(url_for('folder_settings', folder=folder_name))
    old_path = os.path.join(DATA_DIR, folder_name, chapter_name)
    new_path = os.path.join(DATA_DIR, folder_name, new_name)
    if os.path.exists(new_path):
        flash('Name already exists')
    else:
        os.rename(old_path, new_path)
        order = load_order(folder_name)
        if chapter_name in order.get('chapters', []):
            idx = order['chapters'].index(chapter_name)
            order['chapters'][idx] = new_name
            save_order(folder_name, order)
        flash('Chapter renamed')
    return redirect(url_for('folder_settings', folder=folder_name))


@app.route('/folder/<path:folder>/chapter/<chapter>/notes/download')
def download_note(folder, chapter):
    folder_name = sanitize_path(folder)
    chapter_name = safe_name(chapter)
    path = os.path.join(DATA_DIR, folder_name, chapter_name)
    note_name = note_filename(chapter_name)
    return send_from_directory(path, note_name, as_attachment=True, download_name=note_name)


@app.route('/folder/<path:folder>/chapter/<chapter>/chapter.docx')
def download_chapter_docx(folder, chapter):
    folder_name = sanitize_path(folder)
    chapter_name = safe_name(chapter)
    path = os.path.join(DATA_DIR, folder_name, chapter_name)
    book = folder_name.split('/')[0]
    author = read_author(book)
    parts = [book]
    if author:
        parts.append(author)
    parts.append(chapter_name)
    filename = " - ".join(parts) + ".docx"
    return send_from_directory(
        path,
        'chapter.docx',
        as_attachment=True,
        download_name=filename,
    )




@app.route('/folder/<path:folder>/combined.docx')
def download_combined_docx(folder):
    """Download a DOCX containing all chapters in a folder."""
    folder_name = sanitize_path(folder)
    path = os.path.join(DATA_DIR, folder_name)
    if not os.path.isdir(path):
        flash('Folder not found')
        return redirect(url_for('index'))
    chapters = list_chapters(folder_name, include_closed=True)
    if not chapters:
        flash('No chapters to combine')
        return redirect(url_for('view_folder', folder=folder_name))
    doc = Document()
    for idx, chap in enumerate(chapters):
        doc.add_heading(chap, level=1)
        html_file = os.path.join(path, chap, 'chapter.html')
        if os.path.isfile(html_file):
            with open(html_file) as f:
                append_html_to_docx(doc, f.read())
        if idx < len(chapters) - 1:
            doc.add_page_break()
    from io import BytesIO
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    book = folder_name.split('/')[0]
    author = read_author(book)
    parts = [book]
    if author:
        parts.append(author)
    filename = " - ".join(parts) + ".docx"
    return send_file(
        bio,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )


@app.route('/folder/<path:folder>/folder/create', methods=['POST'])
def create_subfolder(folder):
    folder_name = sanitize_path(folder)
    name = safe_name(request.form.get('name', ''))
    if not name:
        flash('Folder name required')
        return redirect(url_for('view_folder', folder=folder_name))
    path = os.path.join(DATA_DIR, folder_name, name)
    os.makedirs(path, exist_ok=True)
    order = load_order(folder_name)
    if name not in order.get('folders', []):
        order.setdefault('folders', []).append(name)
        save_order(folder_name, order)
    return redirect(url_for('view_folder', folder=f"{folder_name}/{name}"))


@app.route('/folder/<path:folder>/rename_subfolder/<sub>', methods=['POST'])
def rename_subfolder(folder, sub):
    folder_name = sanitize_path(folder)
    sub_name = safe_name(sub)
    new_name = safe_name(request.form.get('new_name', ''))
    if not new_name:
        flash('New name required')
        return redirect(url_for('folder_settings', folder=folder_name))
    old_path = os.path.join(DATA_DIR, folder_name, sub_name)
    new_path = os.path.join(DATA_DIR, folder_name, new_name)
    if os.path.exists(new_path):
        flash('Name already exists')
    else:
        os.rename(old_path, new_path)
        order = load_order(folder_name)
        if sub_name in order.get('folders', []):
            idx = order['folders'].index(sub_name)
            order['folders'][idx] = new_name
            save_order(folder_name, order)
        flash('Sub-folder renamed')
    return redirect(url_for('folder_settings', folder=folder_name))


@app.route('/folder/<path:folder>/stats')
def folder_stats(folder):
    folder_name = sanitize_path(folder)
    path = os.path.join(DATA_DIR, folder_name)
    days = int(request.args.get('days', 7))
    total_words = 0
    words_per_day = {}
    for root, dirs, files in os.walk(path):
        if 'chapter.html' in files:
            html_path = os.path.join(root, 'chapter.html')
            with open(html_path) as f:
                text = html_to_text(f.read())
            count = len(text.split())
            total_words += count
            day = datetime.date.fromtimestamp(os.path.getmtime(html_path)).isoformat()
            words_per_day[day] = words_per_day.get(day, 0) + count
    sorted_days = sorted(words_per_day.keys())
    if days > 0:
        cutoff = (datetime.date.today() - datetime.timedelta(days=days-1)).isoformat()
        sorted_days = [d for d in sorted_days if d >= cutoff]
    chart_labels = sorted(sorted_days)
    chart_data = [words_per_day[d] for d in chart_labels]
    folders = list_books()
    return render_template(
        'stats.html',
        folder=folder_name,
        total_words=total_words,
        words_per_day=words_per_day,
        folders=folders,
        chart_labels=chart_labels,
        chart_data=chart_data,
        days=days,
    )


@app.route('/changelog')
def changelog_page():
    """Display the changelog markdown."""
    path = os.path.join(os.path.dirname(__file__), 'CHANGELOG.md')
    content = ''
    if os.path.isfile(path):
        with open(path) as f:
            content = f.read()
    folders = list_books()
    return render_template('changelog.html', content=content, folders=folders)


@app.route('/search')
def search():
    query = request.args.get('q', '').strip()
    results = []
    if query:
        qlower = query.lower()
        for root, dirs, files in os.walk(DATA_DIR):
            rel = os.path.relpath(root, DATA_DIR)
            if 'chapter.html' in files:
                chap = os.path.basename(root)
                with open(os.path.join(root, 'chapter.html')) as f:
                    text = html_to_text(f.read())
                if qlower in text.lower():
                    results.append({'folder': rel, 'chapter': chap, 'type': 'chapter'})
            for fn in files:
                if fn.endswith('_notes.txt'):
                    chap = os.path.basename(root)
                    with open(os.path.join(root, fn)) as nf:
                        text = nf.read()
                    if qlower in text.lower():
                        results.append({'folder': rel, 'chapter': chap, 'type': 'notes'})
    folders = list_books()
    return render_template('search.html', q=query, results=results, folders=folders)


@app.route('/assets/<path:filename>')
def asset_file(filename):
    """Serve files from the assets directory."""
    path = os.path.join(os.path.dirname(__file__), 'assets')
    return send_from_directory(path, filename)


@app.route('/help')
def help_page():
    """Display a basic help page."""
    folders = list_books()
    return render_template('help.html', folders=folders)


@app.route('/download_database')
def download_database():
    """Download the entire data directory as a zip file."""
    from io import BytesIO
    import zipfile

    mem = BytesIO()
    with zipfile.ZipFile(mem, 'w', zipfile.ZIP_DEFLATED) as zf:
        for root, dirs, files in os.walk(DATA_DIR):
            for fname in files:
                path = os.path.join(root, fname)
                rel = os.path.relpath(path, DATA_DIR)
                zf.write(path, rel)
    mem.seek(0)
    return send_file(mem, as_attachment=True, download_name='calwriter_data.zip', mimetype='application/zip')


@app.route('/export_db')
def export_db():
    """Export the database as a .calwdb archive."""
    from io import BytesIO
    import zipfile

    mem = BytesIO()
    with zipfile.ZipFile(mem, 'w', zipfile.ZIP_DEFLATED) as zf:
        for root, dirs, files in os.walk(DATA_DIR):
            for fname in files:
                path = os.path.join(root, fname)
                rel = os.path.relpath(path, DATA_DIR)
                zf.write(path, rel)
        metadata_path = os.path.join(DATA_DIR, 'metadata.json')
        if not os.path.isfile(metadata_path):
            zf.writestr('metadata.json', json.dumps({'version': VERSION}))
    mem.seek(0)
    timestamp = datetime.datetime.now().strftime('%Y%m%d-%H%M%S')
    filename = f"calwriter - {timestamp}.calwdb"
    return send_file(
        mem,
        as_attachment=True,
        download_name=filename,
        mimetype='application/x-calwriter-db',
    )


@app.route('/import_db', methods=['POST'])
def import_db():
    """Import a .calwdb archive into the data directory."""
    file = request.files.get('file')
    if not file or not file.filename.endswith('.calwdb'):
        flash('Invalid file')
        return redirect(url_for('index'))
    import zipfile
    import tempfile
    import shutil

    try:
        with zipfile.ZipFile(file) as zf:
            names = zf.namelist()
            required = {'metadata.json', 'settings.json'}
            missing = [r for r in required if r not in names]
            if missing:
                flash('Archive is missing: ' + ', '.join(missing))
                return redirect(url_for('index'))
            for name in names:
                if name.startswith('/') or '..' in name.split('/'):
                    flash('Invalid path in archive')
                    return redirect(url_for('index'))
            temp_dir = tempfile.mkdtemp()
            zf.extractall(temp_dir)
    except zipfile.BadZipFile:
        flash('File is not a valid archive')
        return redirect(url_for('index'))

    for root, dirs, files in os.walk(temp_dir):
        rel = os.path.relpath(root, temp_dir)
        dest = os.path.join(DATA_DIR, rel) if rel != '.' else DATA_DIR
        os.makedirs(dest, exist_ok=True)
        for fname in files:
            shutil.move(os.path.join(root, fname), os.path.join(dest, fname))
    shutil.rmtree(temp_dir, ignore_errors=True)
    flash('Database imported')
    return redirect(url_for('index'))


@app.route('/about')
def about_page():
    """Show license information."""
    folders = list_books()
    return render_template('about.html', folders=folders)


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000)
